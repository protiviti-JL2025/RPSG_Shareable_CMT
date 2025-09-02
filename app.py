import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime
from streamlit_quill import st_quill
import re

# =========================
# CLAUSES LIBRARIES
# =========================
CLAUSES_NDA = {
    #removing all the serial numbers. 
1: """1. Agreement Terminologies: 
    
In this Agreement except where the context otherwise requires:

(a) In this Agreement the terms listed below shall have the following meanings:

“Affiliate”) means with respect to any Party any Person which, directly or indirectly, (a) Controls such Party, (b) is Controlled by such Party, (c) is Controlled by the same Person who, directly or indirectly, Controls such Party. 

“Agreement” means this Agreement, as may be amended from time to time.
“Applicable Law” means applicable laws, by-laws, rules, regulations, orders, ordinances, protocols, codes, guidelines, policies, notices, directions, judgments, decrees or other requirements or official directive of any Governmental Authority or Person acting under the authority of any Governmental Authority, as applicable to the Parties, as the case may be.

“Confidential Information” includes any written information, content and data shared by the Disclosing Party, whether relating to business idea, operating models, cost plans, customer lists, vendor details, marketing plans, launch plans, know-how, methodology or data, in any written form or medium, tangible or intangible, used in or relating to the assets, properties, business activities, or operations of the Disclosing Party and/or its Affiliates, which is disclosed in writing by the Disclosing Party and/or its representatives to the Receiving Party and/or its representatives, in any written form or media. An illustrative list of Confidential Information is set out below:
a) any information relating to future and proposed plans, processes, products, services and sales, including but not limited to, the information that is shared with the Receiving Party and/or its Representatives or that the Receiving Party or their Representatives come across during the course of this Agreement;

b) any information relating to the business, affairs, practices or methods, services, operational processes, marketing activities, technical know–how, administrative and/or organizational matters relating to the Disclosing Party and/or its Affiliates provided by the Disclosing Party and/or its Affiliates or their Representatives, customer data, vendor details, in particular, names, addresses of its present or prospective investors or target companies/firms;

c) information of a business nature, such as financial statements, marketing plans, business plans, strategies, forecasts, unpublished financial information, budgets, projections, information, and data concerning costs, profits, market share, sales, current or planned distribution methods and processes, target company/firm lists, market studies, business plans, or information regarding, investors or lenders of the Disclosing Party and/or its Affiliates;

d) any other information obtained from the Receiving Party and/or its Representatives in relation to the Disclosing Party, which otherwise fall within the scope of this definition of “Confidential Information.
“Control” means with respect to any Person, means: (a) the possession, directly or indirectly, of the power to direct or cause the direction of the management and policies of such Person whether through the ownership of voting securities, by agreement or otherwise or the power to elect more than one-half of the directors, partners or other individuals exercising similar authority with respect to such Person, or (b) the possession, directly or indirectly, of a voting interest of equal to or more than 50% (Fifty Percent) in a Person.

“Disclosing Party”, shall mean, RPSG Ventures Limited and its officers, agents and other persons engaged by the Company for the purpose of this agreement
“Governmental Authority” shall mean any nation or government or any province, state or any other political subdivision thereof; any entity, authority or body exercising executive, legislative, judicial, regulatory or administrative functions of or pertaining to government, including any government authority, agency, department, board, commission or instrumentality of India or any other country, state or jurisdiction.
"Representative" means, as to any Person, such Person's affiliates and its and their respective directors, officers, members, employees, agents, partners, financing providers, co-investors, investors, consultants, advisors (including, without limitation, financial advisors, counsel and accountants) and controlling Persons.
“Receiving Party”, shall mean ______________ including its  officers, agents and other persons engaged by the Company for the purpose of this agreement.

“Third Party” shall mean any Person other than the Receiving Party or Disclosing Party not being a Party to this Agreement.""",

2: """2. Receiving Party Agreement: The Receiving Party agrees that all written information disclosed by the Disclosing Party Group pursuant to this Agreement shall be considered Confidential Information, unless otherwise specified in writing by the Disclosing Party. By executing this Agreement, the Receiving Party acknowledges that the Disclosing Party derives independent economic value from the Confidential Information not being generally known and that the disclosure of the Confidential Information is subject to the Receiving Party keeping such information in confidence. The Receiving Party shall inform each of its Representatives to whom it intends to disclose Confidential Information in accordance with this Agreement that the obligations imposed on the Disclosing Party in relation to the Confidential Information shall apply in the same manner to such Representatives.""",

3: """3. Non-disclosure Obligations: 

The Receiving Party shall utilize the Confidential Information provided by any Disclosing Party Group member exclusively for the Business Purpose and shall not disclose, publish or disseminate (except with the prior written consent of the Disclosing Party) any Confidential Information to any Third Party other than its Representatives and such other persons whom the Disclosing Party grants its express written consent for disclosure of Confidential Information, who shall be necessarily required to possess such Confidential Information in order for the Receiving Party to fulfil the Business Purpose. Without limiting the generality of the foregoing, the Receiving Party shall, and shall procure and ensure that each of its Representatives to whom the Receiving Party discloses any Confidential Information in accordance with this Agreement shall:

A. Hold such Confidential Information in strict confidence and take commercially reasonable precautions and adequate measures to protect and safeguard the Confidential Information against unauthorized use, publication or disclosure.

B. Not use any of the Confidential Information except in furtherance of the Business Purpose.

C. Not, directly or indirectly, in any way, disclose any of the Confidential Information to any person except as specifically authorized by the Disclosing Party in accordance with this Agreement.

Restrict the access to all Confidential Information by the Representatives on a strictly “need to know” basis for the performance of their duties in furtherance of the Business Purpose.""",

4: """4. Exceptions:  The confidentiality obligations hereunder shall not apply to Confidential Information which (i) is or later becomes public knowledge, except as a result of any unauthorised disclosure by the Receiving Party or its Representatives pursuant to this Agreement; or (ii) was rightfully in possession of the Receiving Party or its Representatives, on a non-confidential basis, prior to its receipt thereof from any Disclosing Party Group member as can be reasonably demonstrated by the Receiving Party via written records, or (iii) is independently developed by the Receiving Party or its Representatives without the use of any Confidential Information as can be reasonably demonstrated by the Receiving Party via written records.""",

5: """5. Return of Confidential Information: Upon the written request of the Disclosing Party or termination of this Agreement, whichever is the earlier, the Receiving Party shall, and shall procure and ensure that each of its Representatives shall, within 14 days thereafter deliver to the Disclosing Party all records, notes, and other written, printed, or tangible materials either in soft or hard copy form which is in the possession of the Receiving Party or its Representatives, embodying or pertaining to the Confidential Information. The Receiving Party shall promptly notify the Disclosing Party following completion of the foregoing obligation. Notwithstanding the foregoing, the Receiving Party and its Representatives (i) may retain copies of the Confidential Information to the extent that such retention is required to demonstrate compliance with applicable law, rule, regulation or professional standards, or to comply with a bona fide document retention policy, provided, however, that any such information so retained shall be held in compliance with the terms of this Agreement for a period of 7 (seven) years from the date of this Agreement and (ii) shall, to the extent that (i) above is inapplicable to Confidential Information that is electronically stored, destroy such electronically stored Confidential Information only.""",

6: """6. Unpublished Price Sensitive Information: The Receiving Party acknowledges that, in connection with and in furtherance of the Proposed Transaction, the Receiving Party may receive Confidential Information which may contain unpublished price sensitive information (UPSI) as defined under the SEBI (Prohibition of Insider Trading) Regulations, 2015, as amended from time to time. Each Party represents that it is aware of the securities laws prevalent in India, including the SEBI (Prohibition of Insider Trading) Regulations, 2015, as amended from time to time and the respective parties shall be responsible for compliance with such laws in respect of receipt and use of UPSI.""",

7: """7. No Representation or Warranty: Except as maybe provided in a definitive agreement between the Parties in connection with the Proposed Transaction, the Disclosing Party does not make any representation or warranty as to the accuracy or completeness of the Confidential Information or of any other information provided, or as to the reasonableness of any assumptions on which any of the same is based, to the Receiving Party or its Representatives, and accordingly, the Receiving Party agrees that the Disclosing Party Group and any of its directors, officers, employees, advisers or agents shall have no liability towards the Receiving Party which may result from the Receiving Party’s unauthorized, use, disclosure or possession of the Confidential Information nor for any claims of Third Parties or as a result of their reliance on any Confidential Information nor for any opinions, projections or forecasts expressed or made by them nor for any errors, omissions or mis-statements made by any of them, and agrees that the Confidential Information is subject to change without notice at any time. In furnishing any Confidential Information no obligation is undertaken by the Disclosing Party to provide any additional information.""",

8: """8. No grant of any right, title, or interest in the Confidential Information: The Confidential Information, including without limitation any patents, copyrights, trademarks, or other intellectual property rights (present or future) in such Confidential Information, shall at all times remain the sole and exclusive property of the Disclosing Party Group. In no situation whatsoever, shall the Receiving Party have any title, right, interest, or claim over such Confidential Information.""",

9: """9. Compelled Disclosure: Pursuant to any Applicable Law, if the Receiving Party or any of its Representatives receives any notice or order by any judicial, Governmental Authority or regulatory entity to disclose any or all Confidential Information, then the Receiving Party shall, and shall procure that such Representatives shall, (to the extent permitted by Applicable Law) make reasonable efforts to promptly notify the Disclosing Party so that the Disclosing Party has the opportunity to intercede and contest such disclosure and the Receiving Party shall, and shall procure that such Representatives shall, wherever reasonably required, cooperate with the Disclosing Party in contesting such a disclosure. The Receiving Party shall, and shall procure that such Representatives, furnish only such part of the Confidential Information that the Receiving Party or such Representatives are legally compelled to disclose to the extent legally permissible.""",

10: """10. No Trade Obligation: Notwithstanding anything contained in this Agreement, the Receiving Party agrees that neither the Receiving Party or its Affiliates shall acquire any interest (whether economic or otherwise) in the Company, other than by way of the Proposed Transaction, for a period of 6 (six) months from the date of this Agreement.""",

11: """11. Losses: The prevailing Party in any dispute between the Parties shall be entitled to recover its reasonable costs and expenses (including their attorney’s fees and costs) in connection with such action.""",

12: """12. Notices: Any notice, request or instruction to be given hereunder by any Party to the other Party shall be in writing, in English language and delivered personally, or sent by registered mail postage prepaid or courier or electronic mail addressed to the concerned Party at the address set forth below or any other address subsequently notified to the other Parties.

Company: 
Address:         RPSG House, 2/4, Judges Court Road,
Kolkata 700027, West Bengal
Attention:      Legal Department 
Email Address:  [•]

For receiving party
Address:        [•]
Attention: Mr.  [•]
Email Address:  [•]

Any notice, request or instruction: (i) sent by email, shall be deemed received when sent; (ii) sent by hand, shall be deemed received when delivered; or (iii) sent by post, shall be deemed received 48 hours after posting.""",

13: """13. Counterparts: This Agreement may be executed in two counterparts, each of which shall be deemed an original, but all of which together shall constitute one and the same instrument.""",

14: """14. Term and Termination:

If either Party decides not to proceed with the Business Purpose with the other Party, it shall notify the other Party in writing immediately (such notice, a “Termination Notice”). This Agreement shall commence on the Execution Date and remain in full effect until earlier of: (a) execution and delivery of the definitive agreements regarding the Proposed Transaction; or (b) 24 months from the Execution Date.""",

15: """15. Remedies: The Receiving Party understands and acknowledges that any disclosure or misappropriation of any of the Confidential Information in violation of this Agreement may cause the Disclosing Party Group irreparable harm, the amount of which may be difficult to ascertain and, therefore, agrees that the Disclosing Party shall have the right to apply to a court of competent jurisdiction for an order restraining any such further disclosure or misappropriation and for such other relief as the Disclosing Party shall deem appropriate. Such right of the Disclosing Party shall be in addition to any other remedies available to the Disclosing Party at law or in equity.""",

16: """16. Non-Disclosure by the Company: Except as required by law, regulation, legal process or any court order, without the Receiving Party prior written consent, the Company and its respective Representatives shall not, directly or indirectly, identify the Receiving Party or its affiliates by name or identifiable description as being involved in discussions or negotiations concerning the Proposed Transaction, or disclose any of the terms, conditions, work product or analysis prepared or submitted by the Receiving Party in connection therewith, to any person other than a Representative of the Company who reasonably requires access to such information in connection with the Proposed Transaction.""",

17: """17. Entire Agreement: This Agreement embodies the entire understanding between the Parties relating to the subject matter of this Agreement and supersedes any and all prior negotiations, correspondence, understandings and agreements between the Parties relating to the subject matter of this Agreement. This Agreement shall not be modified except by a writing duly executed by authorized representatives of all Parties. Should any provision of this Agreement be found unenforceable, such provision or part thereof, to the minimum extent required, shall be deemed to be deleted from this Agreement and the validity and enforceability of the remainder of this Agreement shall still be in effect.""",

18: """18. No Waiver: The failure of the Disclosing Party to require performance by the Receiving Party of any provision of this Agreement shall in no way effect the full right to require such performance at any time thereafter.""",

19: """19. Assignment: The Receiving Party shall have no right to assign or otherwise transfer, in whole or in part, any of its rights or obligations under this Agreement without obtaining prior written consent from the Disclosing Party.""",

20: """20. Third Party Rights: Except as expressly provided in this Agreement, no Third Party shall have any right to enforce any term of this Agreement.""",

21: """21. Governing Law: This Agreement shall be governed by and construed in accordance with the laws of India, without regard to its choice of law provisions and the Courts in Kolkata, India shall have non-exclusive jurisdiction over any dispute hereunder.""",

22: """22. Amendment: This Agreement constitutes the sole understanding of the Parties about this subject matter and may not be amended or modified except in writing signed by each of the Parties to the Agreement."""
}

# 21 CLAUSES for PURCHASE AGREEMENT (grouped exactly as your draft)

CLAUSES_PA = {
    1: """1. DEFINITIONS:

In this AGREEMENT, the following expressions shall have, where the context so admits, the meanings assigned thereto.

(a) "AGREEMENT" shall mean this document together with appendices hereto if any, and shall include any modifications and alterations hereto made in writing.
(b) "Effective Date" shall mean the date on which the authorised representatives of the parties have duly executed this AGREEMENT.
(c) "Licenses" shall mean clearances, licenses, registrations, nominations and permits required under Food Safety and Standards Act 2006, Food Safety and (Standards Packaging and Labelling) Regulations 2011, Legal Metrology Act 2009 and Legal Metrology (Packaging Commodity) Rules 2011 as on today and/or amended from time to time and/or any other law required to run the business of sale of the food articles.
(d) "Products" shall mean and refer to all type of articles manufactured, processed, supplied, distributed and marketed by the Supplier that had been mutually agreed herein or shall be determined by the Company at a future date, to be supplied at the retail outlets of the Company or at any other place that the Company may demarcate from time to time.
(e) "Parties" shall mean the Company or ___________________ together and party shall mean either Spencer’s Retail Limited or _______________, as the case may be.
(f) "Stores" will mean and include all such Spencer’s stores which are in operation or which may be opened during the Term of this Agreement.""",
    2: """2. SCOPE OF WORK

2.1 The Company shall periodically place purchase orders for the products and the Supplier shall deliver the same to the Company in such manner and at such place as required/intimated by the Company from time to time.

2.2 The Company shall inform the Supplier of the specifications of the products and may ask for the samples prior to such supply and the Supplier shall supply the products strictly as per the specifications as shall be prescribed by the Company.

2.3 The Supplier shall ensure that the products are neatly and cleanly packed in such manner and quantity as required by the Company before supply.

The Supplier shall ensure that each and every package containing the materials/products shall contain the following in the main display panel: (i) identity of the commodity; (ii) accurate number of the commodity contained; (iv) Maximum Retail Price; (v) name and address of the manufacturer; (vi) date when packed and serial number; (vii) contact details in case of consumer dispute; (viii) Veg/Non-Veg logo; (ix) List of ingredients; (x) Nutritional Panel; (xi) Batch number.

In the event of any action being brought against the Company or any penalties being imposed on the Company for failure on the part of the Supplier to conform to applicable laws and regulations:
(a) the Supplier shall indemnify the Company for any loss or damage incurred by the Company as a result of the action or penalty imposed;
(b) the Supplier shall indemnify the Company for all costs, expenses and damages incurred by the Company in any and all proceedings, civil or criminal, that may be brought against the Company for failure to conform to the Packaging Rules;
(c) the Company shall have the option of terminating this Agreement within 10 days of knowledge of the action being brought against the Company, provided such breach on part of the Supplier had not been addressed and/or steps initiated for rectification.

2.4 The Company shall provide the Supplier with the schedule of its requirement sufficiently in advance to enable planning. However, the schedule will be tentative. Supplier shall be ready to comply within the time prescribed by the Company, not exceeding 14 (Fourteen) days from receipt of the purchase/supply request.

2.5 The Supplier shall at all times supply the products to the Company in the exact quantity and manner at the place confirmed by the Company.

2.6 The Supplier shall maintain a system of ‘traceability’.

2.7 Order Cancellation
(a) Any time before scheduled commencement of manufacturing of the products as mentioned in the PO, the Company may cancel the order. No compensation will be paid.
(b) Company may reject and refuse to pay for Products which (i) are not manufactured/packed as per specifications; (ii) are damaged upon receipt; (iii) are not in compliance with this Agreement.

2.8 RISK IN TRANSIT: Manufacturer shall be solely responsible for all risk and damages during transit until received by the Company or its designee. Any transit/freight claims shall be handled by the Manufacturer.""",
    3: """3. PRINCIPAL TO PRINCIPAL BASIS

This AGREEMENT is on a Principal to Principal basis.

a. COMPANY shall not be responsible for any acts/omissions of VENDOR’s personnel working under VENDOR’s supervision.
b. COMPANY and VENDOR are independent parties; nothing herein creates JV, partnership, agency, consultancy or employment.""",
    4: """4. PRICE & PAYMENT

4.1 Prices as per annexure; any change to be intimated at least 30 days in advance.

4.2 Price payable = invoice price minus margin/discounts offered by Supplier; inclusive of all taxes, duties, cess and deliveries, borne by Supplier.

4.3 Payment within period specified in Annexure from receipt of goods, unless defects/damages/pilferages/complaints/claims arise within that period.""",
    5: """5. RIGHTS AND DUTIES OF THE SUPPLIER

5.1 Diligently carry out the work; improve quality, reduce cost, improve productivity.
5.2 Carry out supply strictly as per statutory provisions and Company checks.
5.3 Inform Company immediately of any quality problem.
5.4 Ensure products and packaging conform to Legal Metrology and FSSAI, etc.
5.5 Ensure no stoppage/delay from accepted schedule.
5.6 Supply FOR DC or FOR retail outlet as per PO terms.
5.7 If unable to adhere to schedule, inform Company; Company may source outside and debit loss/additional cost (excluding force majeure).
5.8 Supplier responsible for GST/levies, etc.
5.9 Obtain and comply with all sanctions/licenses; indemnify Company for non-compliance.
5.10 No child labour in manufacturing/packaging/distribution.""",
    6: """6. CONSIDERATION

6.1 Supplier has agreed to pay/allow consideration and credit terms as in annexure.""",
    7: """7. RETURN POLICY

7.1 Company may return non-compliant products within 30 days of receipt; transport cost to Supplier’s account.
7.2 For any consumer quality complaint, Supplier shall replace free of cost and indemnify Company for losses/costs arising therefrom.
7.3 Supplier shall take back unsold non-moving stock from the point of delivery.""",
    8: """8. DAMAGE POLICY

8.1 Supplier shall replace expired product if supplied to Company; no liability for mishandling/improper storage at Company premises.
8.2 Supplier will replace all expired/damaged products as per actual. Both parties will work jointly to reduce expiry/damage.""",
    9: """9. PROMOTION PLAN

9.1 Supplier will provide a quarterly promo calendar, with at least one promotion every month per outlet; special promotions during festivals/regions.
9.2 Supplier will offer consumer promotions for each new outlet for a mutually decided period.""",
    10: """10. RIGHTS AND DUTIES OF THE COMPANY

10.1 Provide specifications/formulations/designs sufficiently in advance.
10.2 Supply schedules and POs to be precise; modifications with Supplier’s consent.
10.3 Pay price as mutually agreed.
10.4 Provide cooperation to improve performance.
10.5 Entitled to reject products not complying with standards; Supplier to replace within 7 days of intimation.
10.6 Supplier to take back rejected stocks within 15 days; during this period Company holds at Supplier’s cost/risk; thereafter costs debited to Supplier.
10.7 Company to share necessary information to enable Supplier’s performance.
10.8 Company may terminate with 4 weeks’ notice; immediate termination for inefficiency/non-compliance/misappropriation, etc.""",
    11: """11. COMPLIANCE OF LEGISLATION

11.1 Supplier shall comply with all applicable laws (FSSAI, Legal Metrology, Labour, PF, ESI, etc.).
11.2 Supplier shall obtain all necessary licenses/permissions at its cost.
11.3 All persons engaged by Supplier are Supplier’s employees; Supplier to comply with labour statutes and ensure no employment claim against Company; includes personnel deputed in stores for promotion.""",
    12: """12. WARRANTY

12.1 Supplier ensures quality manufacturing/packaging under strict checks; complies with applicable law.
12.2 Supplier shall replace products if any quality/manufacturing defect found; provide warranty as per applicable law.
12.3 Supplier shall give proper invoice and warrants that:
(a) duly organized and validly existing;
(b) has power/authority to execute and perform;
(c) corporate actions obtained;
(d) financial capacity;
(e) Agreement is legal/valid/binding;
(f) no conflict with charter/documents/law/agreements;
(g) no actions/proceedings with material adverse effect;
(h) no violations/defaults with material adverse effect; complied with law; no significant liabilities;
(i) representations do not omit material facts;
(j) accepts risk of inadequacy/mistake; Company not liable;
(k) Products are genuine, defect-free, with manufacturer warranty (if any), and meet PO/sample specs.""",
    13: """13. CONFIDENTIALITY

13.1 Parties agree all information obtained under this agreement is confidential and will not be used/shared for any other purpose.
13.2 Logos/trademarks/designs not to be used except as envisaged herein or with prior written consent.
13.3 Supplier shall not assign any part of this agreement without Company’s consent; breach enables immediate termination and damages.""",
    14: """14. INDEMNITY

14.1 Supplier shall indemnify Company against all losses/claims/damages/costs arising from acts/omissions/negligence/breach.
14.2 Supplier specifically indemnifies against claims/fines/penalties arising out of manufacturing defects and/or non-compliance of food laws.
14.3 Supplier shall bear fines/legal fees if required by Company.
14.4 On receipt of any legal notice etc., Company may defend/settle at Supplier’s cost.
14.5 Supplier shall assist in defending; compromise/settlement with mutual agreement.""",
    15: """15. FORCE MAJEURE

15.1 Failure/delay due solely to Force Majeure (act of God, government, riots, war, strikes, lockouts, transport accidents, etc.) shall not be breach, provided the affected party did not cause it, used diligence to avoid/ameliorate, and continues efforts to comply.
15.2 Party suffering Force Majeure shall notify within 7 days and use best efforts to remove/remedy.
15.3 If Force Majeure persists for more than three consecutive months, other party may terminate without liability.""",
    16: """16. TERM AND TERMINATION

16.1 Unless terminated, term is 5 (Five) years from signing.
16.2 Company may terminate without cause by 30 days’ notice.
16.3 Either party may terminate by one month’s notice without reason.

If Supplier cannot provide materials/products within required timeframe, Company may terminate forthwith; Supplier liable for damages due to non-compliance/breach.

Additional grounds for Company’s immediate termination include:
a. Supplier’s failure to perform (save Force Majeure or Company’s sole fault);
b. arrears/unpaid amounts due to Company;
c. false/misleading representations/warranties;
d. illegal/prohibited activities by Supplier or its employees;
e. bankruptcy/insolvency of Supplier;
f. resolution for voluntary winding up;
g. repeated customer grievances on quality;
h. failure to deliver within time;
i. breach of statutory regulations/gross negligence;
j. delivery not matching request, short/excess supply;
k. failure to receive back products despite intimation or after display period.

Additionally, immediate termination if Supplier fails in: (i) Quality; (ii) Service; (iii) Marketing Activities; (iv) Delay in Delivery; (v) Non-acceptance of defective/damaged materials.

After notice period expiry, parties will reconcile accounts within 15 days.""",
    17: """17. QUALITY POLICY

17.1 Supplier shall adhere to standards under Food Laws or other rules; ensure product safety and quality.
17.2 If entire batch is defective, Company may remove from display immediately; cost of goods/logistics debited to Supplier. Supplier to respond within a week with counter-measures, redress complaint and reconcile with aggrieved customer.""",
    18: """18. NOTICES

Any notice/letter required shall be deemed properly served if delivered by RPAD/courier/hand at the addresses mentioned hereinabove.""",
    19: """19. AMENDMENTS & MODIFICATIONS

Any amendment/modification may be initiated by letter from either party; becomes binding when acknowledged/accepted by the other party.""",
    20: """20. INTELLECTUAL PROPERTY RIGHTS

No party shall use the trademark/copyright/IP of the other in any unauthorized manner or in any advertisement/publicity without written permission of the owner.""",
    21: """21. JURISDICTION

The courts of Kolkata shall have jurisdiction over disputes arising out of this agreement."""
}

# ------------------------------------------------
# 1) TEMPLATES
# ------------------------------------------------
NDA_TEMPLATE = """NON-DISCLOSURE AGREEMENT

THIS NON-DISCLOSURE AGREEMENT (this “Agreement”) is entered on {DAY} day of {MONTH} 2025 (“Execution Date”) by and between:

RPSG Ventures Limited, a company incorporated under the laws of India with CIN no. {RPSG_CIN_No}(hereinafter referred to as the “Company/ Disclosing Party”), having a registered office at CESC House, Chowirnghee Square, Kolkata - 700001, West Bengal India (which expression shall, unless repugnant to the context thereof, mean and include its subsidiaries, partners, associates, legal representatives, successors, and permitted assigns);

AND

{Vendor_Name}, a company incorporated under the laws of India with CIN no. {Vendor_CIN_No} and having its registered office at {Vendor_Office}, India. (hereinafter referred to as the “Receiving Party”), (which expression shall, unless repugnant to the context thereof, mean and include its subsidiaries, partners, associates, legal representatives, successors, and permitted assigns).

(both are collectively  referred to as “the Parties”)

WHEREAS:

The Parties are negotiating a possible business transaction referred to below (hereinafter called the “Proposed Transaction”). To facilitate the Proposed Transaction and to evaluate and consider entering into the Transaction, the Disclosing Party shall provide to the Receiving Party the Confidential Information relating to the Company.

The Parties desires to protect its rights and the confidentiality of the Information (as hereinafter defined) and the Parties desire to have access to the Information of the others and set out the terms and conditions to be followed by the Receiving Party with respect to the confidential information.

NOW THEREFORE it is agreed as follows:
"""

PURCHASE_TEMPLATE = """AGREEMENT

THIS AGREEMENT is made and executed at Kolkata on …………………………………………, 202….. between 

M/S. Spencers Retail Limited (formerly known as RP-SG Retail Limited), a Company incorporated under the provisions of the Companies Act, 2013, and having its registered office at Duncan House,   No. 31 Netaji Subhas Road, Kolkata–700001 and Corporate Office at RPSG House, 2/4 Judges Court Road,  Kolkata- 700027 (hereinafter referred to as ‘the Company’ and which term shall, unless repugnant to the context, mean and include all its successors-in-interest and assigns) of the First Part through it Mr…………………, 
AND

<name> (PAN: …………….) son/wife of …………….. aged about ____ years resident of <complete address with Holding No………, Police Station: …………… Post Office: …….., PIN> operating his/her sole proprietorship Business as “………………..”

Or

M/s. …………………, a Partnership Firm, registered under the provisions of the Indian Partnership Act, 1932 bearing registration No. …………… dated …………. (if any), having PAN …………… having its principal place of business at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN ……..…., District: ……………., State: ……….………, represented herein through its Partners (1) Mr./Ms. ………………… S/D/W/o ………… by faith ……..……, by occupation …………, residing at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN …….., District: …………, State: …………> and (2) Mr./Ms. ………………… S/D/W/o ………by faith ………, by occupation …………, residing at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN …….., District: ………, State: …………> duly authorised in this regard by all the other partners vide authorisation letter/certificate dated ………………

Or

………………… LLP, a Limited Liability Partnership, incorporated under the provisions of the Limited Liability Partnership Act, 2008 bearing LLPIN …………., having PAN …………… having its registered office at <complete address with Holding No………, Po-lice Station: …………… Post Office: ………, PIN ……..…., District: ……………., State: ……….………, represented herein through its Designated Partners (1) Mr./Ms. ………………… S/D/W/o ………… by faith ……..……, by occupation …………, residing at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN …….., District: …………, State: …………> having DPIN: ………… and (2) Mr./Ms. ………………… S/D/W/o ………by faith ………, by occupation …………, resid-ing at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN …….., District: ………, State: …………> having DPIN: ………… duly au-thorised in this regard by all the other partners vide authorisation letter/certificate dated ………………

Or

………………… Limited OR Private Limited (CIN: ……………………), a Company incorporated / existing under the provisions of the Companies Act, 2013, having PAN ……………. having its Registered Office at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN ……….., District: …, State: …………> and Corporate Office at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN ……….., District: …, State: …………> represented herein through its authorized signatory Mr./Ms. ………………… <name and designation> S/D/o ………by faith ………, by occupation …………, residing at <complete address with Holding No………, Police Station: …………… Post Office: ………, PIN …….., District: …, State: …………> duly authorised in this regard vide Board resolution dated   ………………

(herein after referred to as ‘the Supplier’ and which term shall, unless repugnant to the context, mean and include its successors and permitted assigns) of the Second Part

WHEREAS:

A. The Supplier is engaged inter-alia in the business of manufacture, sale, marketing and distribution of “___________________ products” (herein after referred to as the “products”) under ‘__________________’ brand.   

B. The Company is engaged in the business of operating retail stores under the various formats, in India under the brand name “Spencer’s”.

C. The Company proposes to sell through its outlets and otherwise, the products of the Supplier and such other items as may be decided mutually by the parties from time to time.

D. The Company has offered the Supplier to supply the products and the Supplier has accepted the same under the following terms and conditions.

NOW THEREFORE THIS AGREEMENT WITNESSETH AS FOLLOWS:
"""

# ------------------------------------------------
# Helpers: highlights and minimal HTML-ish formatting
# ------------------------------------------------

#this function was made so that if duplicate columns are present in the annexure table made using word mode, then the code doesn't crash or throw any errors. (since by default streamlit pyarrow enforces unique column names.)
def render_table(header, body):
    table_html = "<table style='border-collapse: collapse; width: 100%;'>"
    if header:
        table_html += "<tr>" + "".join([f"<th style='border: 1px solid #ddd; padding: 6px;'>{h}</th>" for h in header]) + "</tr>"
    for row in body:
        table_html += "<tr>" + "".join([f"<td style='border: 1px solid #ddd; padding: 6px;'>{c}</td>" for c in row]) + "</tr>"
    table_html += "</table>"
    st.markdown(table_html, unsafe_allow_html=True)


def add_heading(doc: Document, text: str, size: int = 16, center=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _regex_parts_for_terms(terms):
    """Build a regex that finds the longest non-empty terms safely."""
    clean = [t for t in (terms or []) if isinstance(t, str) and t.strip()]
    # sort by length desc to avoid partial matches eating bigger ones
    clean = sorted(set(clean), key=len, reverse=True)
    if not clean:
        return None
    escaped = [re.escape(t) for t in clean]
    return re.compile("(" + "|".join(escaped) + ")", flags=re.IGNORECASE)

def highlight_placeholders(text, mapping):
    """Highlight only values that replaced placeholders."""
    if not text:
        return text
    for key, val in mapping.items():
        if val and isinstance(val, str) and val.strip():
            text = text.replace(val, f"<u><mark>{val}</mark></u>")
    return text

def write_with_highlight(paragraph, text: str, highlight_terms):
    """
    Write text into a docx paragraph. Any appearance of terms is
    underlined + highlighted. Other text is normal.
    """
    if not text:
        return
    if not highlight_terms:
        paragraph.add_run(text)
        return
    rx = _regex_parts_for_terms(highlight_terms)
    if not rx:
        paragraph.add_run(text)
        return

    pos = 0
    for m in rx.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        r = paragraph.add_run(m.group(0))
        r.underline = True
        r.font.highlight_color = 7  # 7 = yellow
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def htmlish_to_blocks(html: str) -> list:
    """
    Convert a small subset of Quill HTML into a list of blocks with structure:
    [{"type": "p"|"ul"|"ol", "items":[{"text":"...", "level":0}, ...]}] or {"type":"p","text":...}
    Supports <p>, <br>, <strong>/<b>, <em>/<i> (lost as plain text for DOCX),
    <ul><li>, <ol type="a|1|i"><li>. Nested lists become level>0.
    """
    if not html:
        return []

    s = html
    # normalize line breaks
    s = s.replace("\r", "")
    s = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    # strip spans but keep inner text
    s = re.sub(r"</?span[^>]*>", "", s, flags=re.I)
    # extract lists
    blocks = []

    # naive splitting on major list/p blocks
    tokens = re.split(r"(?i)(</?ul>|</?ol[^>]*>|</?p>)", s)

    mode = None   # 'ul', 'ol', 'p'
    list_type = "1"  # for <ol> type; default numeric
    pending_items = []

    def flush_list():
        nonlocal pending_items, list_type
        if pending_items:
            blocks.append({"type": mode, "list_type": list_type, "items": pending_items[:]})
            pending_items = []

    for tok in tokens:
        if tok is None or tok == "":
            continue
        tag = tok.strip().lower()

        # opening tags
        if tag.startswith("<ul"):
            # flush previous
            if mode == "ol":
                flush_list()
            mode = "ul"
            pending_items = []
            continue
        if tag.startswith("</ul"):
            if mode == "ul":
                flush_list()
            mode = None
            continue

        if tag.startswith("<ol"):
            # detect type
            m = re.search(r'type\s*=\s*"([aAiI1])"', tok)
            list_type = m.group(1) if m else "1"
            if mode == "ul":
                flush_list()
            mode = "ol"
            pending_items = []
            continue
        if tag.startswith("</ol"):
            if mode == "ol":
                flush_list()
            mode = None
            list_type = "1"
            continue

        if tag == "<p>":
            if mode in ("ul", "ol"):
                # treat paragraph inside list as item with level 0
                continue
            mode = "p"
            continue
        if tag == "</p>":
            mode = None
            continue

        # content
        content = tok

        # <li> items
        li_splits = re.split(r"(?i)</?li[^>]*>", content)
        if len(li_splits) > 1 and mode in ("ul", "ol"):
            # build items from inner text
            for piece in li_splits:
                piece = piece.strip()
                if not piece:
                    continue
                # naive nesting level by counting <ul> / <ol> inside (rare with Quill blocks)
                level = len(re.findall(r"(?i)<ul|<ol", piece))
                # strip remaining tags
                txt = re.sub(r"<[^>]+>", "", piece).strip()
                if txt:
                    pending_items.append({"text": txt, "level": level})
            continue

        # a raw text while in paragraph
        if mode == "p":
            txt = re.sub(r"<[^>]+>", "", content).strip()
            if txt:
                blocks.append({"type": "p", "text": txt})
            continue

        # outside any tag: plain text
        fallback = re.sub(r"<[^>]+>", "", content).strip()
        if fallback:
            blocks.append({"type": "p", "text": fallback})

    return blocks

def add_blocks_to_doc(doc: Document, html_text: str, highlight_terms):
    """
    Add converted blocks into docx with numbering/bullets and indent.
    """
    blocks = htmlish_to_blocks(html_text or "")
    if not blocks:
        return

    # counters for ordered lists per level
    ol_counters = {}

    for b in blocks:
        if b["type"] == "p":
            p = doc.add_paragraph()
            write_with_highlight(p, b["text"], highlight_terms)
            continue

        if b["type"] in ("ul", "ol"):
            list_type = b.get("list_type", "1")
            ol_counters = {}

            for it in b.get("items", []):
                level = min(int(it.get("level", 0)), 3)
                p = doc.add_paragraph()
                # indent via left indentation (approx 0.5" per level)
                p.paragraph_format.left_indent = Inches(0.5 * level)

                # marker
                if b["type"] == "ul":
                    marker = "• "
                else:
                    # ordered list
                    idx = ol_counters.get(level, 0) + 1
                    ol_counters[level] = idx
                    if list_type in ("a", "A"):
                        base = chr(ord('a') + (idx - 1))
                        marker = f"({base if list_type=='a' else base.upper()}) "
                    elif list_type in ("i", "I"):
                        # simple roman up to 12
                        romans = ["i","ii","iii","iv","v","vi","vii","viii","ix","x","xi","xii"]
                        rm = romans[idx-1] if 1 <= idx <= len(romans) else str(idx)
                        marker = f"{rm if list_type=='i' else rm.upper()}. "
                    else:
                        marker = f"{idx}. "

                run_m = p.add_run(marker)
                run_m.bold = False
                write_with_highlight(p, it["text"], highlight_terms)

# ------------------------------------------------
# DOCX builder with sections + clause titles + highlighted inputs
# ------------------------------------------------
def extract_title_from_clause(text: str) -> str:
    """
    Extract a short title for dropdowns:
    - If clause starts with 'N. TITLE: ...' or 'TITLE:' → use before colon.
    - Else → return everything up to the first full stop (.)
    """
    t = re.sub(r"<[^>]+>", "", (text or "")).strip()
    if not t:
        return "Clause"

    first_line = t.splitlines()[0]

    # If contains colon, take before colon
    if ":" in first_line:
        head = first_line.split(":", 1)[0].strip()
        return head

    # If starts with number and a dot, strip the number
    m = re.match(r"^\s*\d+\.\s*(.+)$", first_line)
    if m:
        first_line = m.group(1).strip()

    # Stop at first full stop if present
    if "." in first_line:
        return first_line.split(".", 1)[0].strip()

    # Fallback: whole line
    return first_line.strip()

#so now the NDA dropdown text in the box will only have some words instead of the whole clause. 

def build_docx(
    preamble_html: str,
    clauses_dict: dict[int, str],
    custom_clauses: list[dict],
    parties: list[str],
    annexures: list[dict],
    title="AGREEMENT",
    highlight_terms: list[str] | None = None
) -> BytesIO:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Title
    add_heading(doc, title.upper(), 18, True)
    # Preamble
    add_blocks_to_doc(doc, preamble_html, highlight_terms)
    doc.add_paragraph("")

    # Sections (standard clauses)
    for key in sorted(clauses_dict.keys()):
        cl_html = clauses_dict[key]
        # Skip this logic for the body, just add the raw blocks
        add_blocks_to_doc(doc, cl_html, highlight_terms)
        doc.add_paragraph("")

    # Custom clauses (if any)
    if custom_clauses:
        start_num = len(clauses_dict) + 1
        for idx, c in enumerate(custom_clauses, start=start_num):
            t = f"{idx}. {c.get('title', 'Custom Clause')}"
            p = doc.add_paragraph()
            rr = p.add_run(t)
            rr.bold = True
            rr.font.size = Pt(12)
            add_blocks_to_doc(doc, c.get("text",""), highlight_terms)
            doc.add_paragraph("")

    # Parties
    if parties:
        doc.add_paragraph()
        add_heading(doc, "PARTIES", 14, False)
        for i, ptxt in enumerate(parties, 1):
            p = doc.add_paragraph()
            write_with_highlight(p, f"{i}. {ptxt}", highlight_terms)

    # Multiple Annexures
    if annexures:
        for i, annex in enumerate(annexures):
            doc.add_page_break()
            annex_title = annex.get("title", f"ANNEXURE {i+1}")
            add_heading(doc, annex_title.upper(), 14, True)

            # Annexure note
            annex_note = annex.get("note")
            if annex_note:
                p = doc.add_paragraph()
                write_with_highlight(p, annex_note, highlight_terms)
                doc.add_paragraph("")  # spacing

            # Annexure Excel data
            annexure_excel = annex.get("excel_data")
            df = None
            if annexure_excel is not None:
                try:
                    annexure_excel.seek(0)
                    df = pd.read_excel(annexure_excel)
                except Exception as e:
                    df = None
                    doc.add_paragraph(f"[{annex_title} (Excel) could not be read: {e}]")
            if df is not None:
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for j, c in enumerate(df.columns):
                    hdr[j].text = str(c)
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for j, val in enumerate(row):
                        row_cells[j].text = "" if pd.isna(val) else str(val)

            # Annexure Word table data
            annexure_word_rows = annex.get("word_data")
            if annexure_word_rows:
                if df is not None:
                    doc.add_paragraph("")  # spacing between tables
                maxc = max((len(r) for r in annexure_word_rows), default=0)
                if maxc > 0:
                    table2 = doc.add_table(rows=0, cols=maxc)
                    table2.style = "Table Grid"
                    for r in annexure_word_rows:
                        cells = table2.add_row().cells
                        for j in range(maxc):
                            cells[j].text = (r[j] if j < len(r) else "") or ""

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ------------------------------------------------
# Streamlit App
# ------------------------------------------------
st.set_page_config(page_title="Interactive Contract Generator", page_icon="📄", layout="wide")
st.title("📄 Interactive Contract Generator")

# Top-level role toggle
if "role" not in st.session_state:
    st.session_state.role = None  # or set "admin" if you want to default to Admin

btn_admin, btn_business = st.columns(2)
if btn_admin.button("👤 ADMIN Users", type="primary", use_container_width=True, key="btn_admin"):
    st.session_state.role = "admin"
    st.rerun()

if btn_business.button("🏢 BUSINESS Users", use_container_width=True, key="btn_business"):
    st.session_state.role = "business"
    st.rerun()

# Gate the rest of the app
if st.session_state.role == "business":
    # blank inside (only the two buttons at the top remain visible)
    st.write("")  # keep it truly blank for now
    st.stop()

# If role is None or "admin", the app proceeds as usual below…
# ==== END OF INSERT ====

# GLOBAL STATE
if "workflow" not in st.session_state:
    st.session_state.workflow = "form"  # form → clauses → annexure → preview
if "answers" not in st.session_state:
    st.session_state.answers = {}
if "clause_store" not in st.session_state:
    st.session_state.clause_store = {}
if "custom_clauses" not in st.session_state:
    st.session_state.custom_clauses = []  # {"title","text"}
if "parties" not in st.session_state:
    st.session_state.parties = []
if "annexures" not in st.session_state:
    st.session_state.annexures = []  # List of dicts: {"title", "note", "mode", "excel_data", "word_data"}
if "prev_contract_type" not in st.session_state:
    st.session_state.prev_contract_type = None
if "last_contract_type" not in st.session_state:
    st.session_state.last_contract_type = ""

# selectors
entity = st.selectbox("Entity", ["", "RPSG Ventures", "PCBL", "Spencers Retail", "AquaPharm", "Nature's Basket"])
contract_type = st.selectbox(
    "Contract Type",
    ["", "NDA", "PURCHASE AGREEMENT", "Vendor Agreement", "Service Agreement", "Marketing Agreement",
     "Software License Agreement", "Supply Agreement", "Lease/Rent Agreement", "Leave and License Agreement",
     "Franchise Agreement", "Distribution Agreement", "Manufacturing Agreement", "Employment Agreement",
     "Collaboration Agreement", "MoU/LoI"],
    index=0
)

# Reset/Sync on contract type change
def sync_clause_store():
    if contract_type == "NDA":
        st.session_state.clause_store = {i: CLAUSES_NDA[i] for i in CLAUSES_NDA}
        st.session_state.custom_clauses = []
    elif contract_type == "PURCHASE AGREEMENT":
        st.session_state.clause_store = {i: CLAUSES_PA[i] for i in CLAUSES_PA}
    else:
        st.session_state.clause_store = {}
        st.session_state.custom_clauses = []

if contract_type != st.session_state.last_contract_type:
    sync_clause_store()
    st.session_state.parties = []
    st.session_state.annexures = [] # Reset annexures
    st.session_state.last_contract_type = contract_type

st.markdown("---")

# ================ STEP 1: FORM UI (collect inputs) ================
if st.session_state.workflow == "form":
    if contract_type == "NDA":
        st.subheader("NDA – Collect Inputs")
        fields = [
            ("DAY", "Signing day (e.g., 21)"),
            ("MONTH", "Signing month (e.g., August)"),
            ("RPSG_CIN_No", "RPSG CIN Number"),
            ("Vendor_Name", "Vendor Company Name"),
            ("Vendor_CIN_No", "Vendor CIN Number"),
            ("Vendor_Office", "Vendor Registered Office Address"),
            ("RPSG_Email", "RPSG Email"),
            ("Vendor_SPOC", "Vendor SPOC Name"),
            ("Vendor_Email", "Vendor Email"),
            ("RPSG_SPOC", "RPSG SPOC Name"),
        ]
        cols = st.columns(2)
        for i, (key, label) in enumerate(fields):
            st.session_state.answers[key] = cols[i % 2].text_input(label, value=st.session_state.answers.get(key, ""))

    elif contract_type == "PURCHASE AGREEMENT":
        st.subheader("Purchase Agreement – Collect Inputs")
        pa_fields = [
            ("Effective_Date", "Effective Date (DD-MM-YYYY)"),
            ("Buyer_Name", "Buyer Name"),
            ("Buyer_Address", "Buyer Address"),
            ("Seller_Name", "Seller Name"),
            ("Seller_Address", "Seller Address"),
            ("Goods_Description", "Goods/Services Description"),
            ("Qty", "Quantity/Scope"),
            ("Unit", "Unit of Measure"),
            ("Unit_Price", "Unit Price / Rate"),
            ("Payment_Terms", "Payment Terms"),
            ("Delivery_Terms", "Delivery/Incoterms"),
            ("Delivery_Schedule", "Delivery Schedule"),
            ("Inspection", "Inspection/Acceptance"),
            ("Warranty", "Warranty"),
            ("Limitation", "Limitation of Liability"),
            ("Termination", "Termination"),
            ("Governing_Law", "Governing Law & Jurisdiction"),
            ("Notices_Email", "Notices Email (Both Parties)"),
        ]
        cols = st.columns(2)
        for i, (key, label) in enumerate(pa_fields):
            st.session_state.answers[key] = cols[i % 2].text_input(label, value=st.session_state.answers.get(key, ""))

    else:
        if contract_type:
            st.info("Proceed to Clauses to add your own sections.")
        else:
            st.info("Select an Entity and a Contract Type to begin.")

    left, right = st.columns([1, 1])
    if left.button("Save & Go to Clauses ➡️", use_container_width=True):
        st.session_state.workflow = "clauses"
        st.rerun()
    right.button("Reset Form", on_click=lambda: st.session_state.answers.clear(), use_container_width=True)

# ======================= STEP 2: SECTIONED CLAUSE EDITOR =======================
elif st.session_state.workflow == "clauses":
    st.subheader("Sections – Edit & Save")
    st.caption("Each section has a rich editor. Use lists and sub-lists; they’ll export as numbered/bulleted items with indentation.")

    # Standard sections first (NDA/PA)
    if contract_type in ("NDA", "PURCHASE AGREEMENT") and st.session_state.clause_store:
        st.markdown("### Standard Sections")
        for key in sorted(st.session_state.clause_store.keys()):
            with st.expander(f"Section: {extract_title_from_clause(st.session_state.clause_store[key])}", expanded=False):
                current = st.session_state.clause_store[key]
                edited = st_quill(value=current, placeholder="Edit section text…", key=f"std_sec_{key}")
                c1, c2 = st.columns([1,1])
                if c1.button("💾 Save", key=f"save_std_{key}", use_container_width=True):
                    if edited is not None:
                        st.session_state.clause_store[key] = edited
                        st.success("Saved.")
                if c2.button("↩️ Reset to library", key=f"reset_std_{key}", use_container_width=True):
                    # reload from library
                    if contract_type == "NDA":
                        st.session_state.clause_store[key] = CLAUSES_NDA[key]
                    else:
                        st.session_state.clause_store[key] = CLAUSES_PA[key]
                    st.info("Reset.")

    # Custom sections for all types
    st.markdown("### ➕ Add Custom Section")
    cc_title = st.text_input("Section Title", key="cc_title")
    cc_body = st_quill(value="", placeholder="Write your custom section…", key="cc_body")
    add_col1, add_col2 = st.columns([1,1])
    if add_col1.button("➕ Add Section", use_container_width=True):
        if cc_title.strip() and cc_body and cc_body.strip():
            st.session_state.custom_clauses.append({"title": cc_title.strip(), "text": cc_body})
            st.success("Custom section added.")
    if add_col2.button("🧹 Clear", use_container_width=True):
        st.session_state.cc_title = ""
        st.session_state.cc_body = ""

    if st.session_state.custom_clauses:
        st.markdown("#### Custom Sections")
        for idx, c in enumerate(st.session_state.custom_clauses, 1):
            with st.expander(f"{idx}. {c.get('title','Custom Section')}", expanded=False):
                new_title = st.text_input("Title", value=c.get("title",""), key=f"cst_t_{idx}")
                new_body = st_quill(value=c.get("text",""), key=f"cst_b_{idx}")
                c1, c2 = st.columns([1,1])
                if c1.button("💾 Save", key=f"save_cst_{idx}", use_container_width=True):
                    c["title"] = new_title
                    c["text"] = new_body
                    st.success("Saved.")
                if c2.button("❌ Remove", key=f"rm_cst_{idx}", use_container_width=True):
                    st.session_state.custom_clauses.pop(idx-1)
                    st.rerun()

    # Parties
    st.markdown("---")
    st.markdown("### 👥 Parties")
    new_party = st.text_input("Add Party (Name, Role, Address/email optional)", key="party_input")
    cpa, cpb = st.columns(2)
    if cpa.button("Add Party"):
        if new_party.strip():
            st.session_state.parties.append(new_party.strip())
    if cpb.button("Clear Parties"):
        st.session_state.parties = []
    if st.session_state.parties:
        st.write("- " + "\n- ".join(st.session_state.parties))

    st.markdown("---")
    nav1, nav2 = st.columns([1,1])
    if nav1.button("⬅️ Back to Form"):
        st.session_state.workflow = "form"
        st.rerun()
    if nav2.button("Save & Go to Annexure ➡️"):
        #st.session_state.workflow = "annexure" if contract_type != "NDA" else "preview"
        st.session_state.workflow = "annexure"
        #so now annexure functionality will also be there for NDA.
        st.rerun()

# ========================= STEP 3: ANNEXURE (Excel + Word modes) =========================
elif st.session_state.workflow == "annexure":
    st.subheader("Annexures")
    st.caption("Add, remove, and edit all annexures for the agreement. Each can be an Excel or Word-style table.")

    if st.button("➕ Add New Annexure", use_container_width=True):
        new_annex_num = len(st.session_state.annexures) + 1
        annex_letter = chr(ord('A') + new_annex_num - 1)
        st.session_state.annexures.append({
            "title": f"Annexure {annex_letter}",
            "note": "",
            "mode": "Word Table Mode",
            "excel_data": None,
            "word_data": [["S.No", "Item/Deliverable", "Description", "Qty", "Unit Price", "Total", "Remarks"]]
        })
        st.rerun()

    st.markdown("---")

    # Loop through and display editors for each annexure
    for i, annex in enumerate(st.session_state.annexures):
        with st.expander(f"**{annex.get('title', f'Annexure {i+1}')}**", expanded=True):
            # Edit title
            new_title = st.text_input("Annexure Title", value=annex["title"], key=f"title_{i}")
            st.session_state.annexures[i]["title"] = new_title

            # Edit mode
            mode = st.radio(
                "Editing Mode",
                ["Excel Mode", "Word Table Mode"],
                index=0 if annex["mode"] == "Excel Mode" else 1,
                key=f"mode_{i}",
                horizontal=True
            )
            st.session_state.annexures[i]["mode"] = mode

            # Excel Mode UI
            if mode == "Excel Mode":
                file = st.file_uploader("Upload Annexure (.xlsx)", type=["xlsx"], key=f"uploader_{i}")
                if file:
                    st.session_state.annexures[i]["excel_data"] = BytesIO(file.getvalue())
                    st.success(f"Uploaded: {file.name} for {annex['title']}")

                st.markdown("**Inline Editor (optional)**")
                raw_df = None
                if st.session_state.annexures[i]["excel_data"]:
                    try:
                        st.session_state.annexures[i]["excel_data"].seek(0)
                        raw_df = pd.read_excel(st.session_state.annexures[i]["excel_data"])
                    except Exception as e:
                        st.error(f"Could not read the uploaded Excel file. Please re-upload. Error: {e}")

                if raw_df is None:
                    raw_df = pd.DataFrame(columns=["S.No","Item/Deliverable","Description","Quantity","Unit Price","Total","Remarks"])

                # Ensure all cells are editable as strings
                raw_df = raw_df.astype(str).replace("nan", "")

                edited_df = st.data_editor(
                    raw_df,
                    num_rows="dynamic",
                    use_container_width=True,
                    key=f"editor_{i}"
                )

                c1, c2 = st.columns(2)
                if c1.button("💾 Save Edits to Excel Buffer", key=f"save_excel_{i}"):
                    out = BytesIO()
                    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
                        edited_df.replace("", pd.NA).to_excel(writer, index=False, sheet_name=annex.get("title", f"Annexure_{i+1}"))
                    out.seek(0)
                    st.session_state.annexures[i]["excel_data"] = out
                    st.success(f"Excel data for {annex['title']} saved.")
                if c2.button("🧹 Clear Excel Data", key=f"clear_excel_{i}"):
                    st.session_state.annexures[i]["excel_data"] = None
                    st.info(f"Cleared Excel data for {annex['title']}.")

           # Word Table Mode UI
            else:
                st.markdown("**Edit Word-style Annexure Table**")
                data = st.session_state.annexures[i]["word_data"]

                import pandas as pd
                df = pd.DataFrame(data)

                # If table is empty, initialize with 2 columns
                if df.empty:
                    df = pd.DataFrame([["", ""]], columns=["Column 1", "Column 2"])

                edited_df = st.data_editor(
                    df,
                    num_rows="dynamic",   # allows adding/removing rows
                    use_container_width=True,
                    hide_index=True,
                    key=f"annex_table_{i}_{id(annex)}"
                )

                # --- Column Controls ---
                col_ctrls = st.columns(2)
                if col_ctrls[0].button("➕ Add Column", key=f"addcol_{i}_{id(annex)}"):
                    new_col_name = f"Column {len(edited_df.columns) + 1}"
                    edited_df[new_col_name] = ""   # add new empty column
                    st.session_state.annexures[i]["word_data"] = edited_df.values.tolist()
                    st.rerun()

                if col_ctrls[1].button("❌ Delete Last Column", key=f"delcol_{i}_{id(annex)}"):
                    if len(edited_df.columns) > 1:   # keep at least one column
                        edited_df = edited_df.iloc[:, :-1]  # drop last column
                        st.session_state.annexures[i]["word_data"] = edited_df.values.tolist()
                        st.rerun()

                # Save back to session state (list of lists)
                st.session_state.annexures[i]["word_data"] = edited_df.values.tolist()

            # Annexure Note
            note = st.text_area(
                "Annexure Note (appears above the table in the contract)",
                value=annex["note"],
                key=f"note_box_{i}_{id(annex)}",
                height=100
            )
            st.session_state.annexures[i]["note"] = note

            # Remove annexure button
            if st.button(f"🗑️ Remove {annex['title']}", key=f"remove_annex_{i}_{id(annex)}"):
                st.session_state.annexures.pop(i)
                st.rerun()

            st.markdown("---")



    nav1, nav2 = st.columns([1, 1])
    if nav1.button("⬅️ Back to Sections"):
        st.session_state.workflow = "clauses"
        st.rerun()
    if nav2.button("Save & Go to Preview ➡️"):
        st.session_state.workflow = "preview"
        st.rerun()

# ========================= STEP 4: PREVIEW & EXPORT =========================
elif st.session_state.workflow == "preview":
    st.subheader("Preview, Download, Submit")

    # Build preamble with the current answers
    #highlight_terms = [v for v in st.session_state.answers.values() if isinstance(v, str) and v.strip()]
     
     # Build placeholder mapping only for defined placeholders
    if contract_type == "NDA":
        placeholder_mapping = {
            "DAY": st.session_state.answers.get("DAY", ""),
            "MONTH": st.session_state.answers.get("MONTH", ""),
            "RPSG_CIN_No": st.session_state.answers.get("RPSG_CIN_No", ""),
            "Vendor_Name": st.session_state.answers.get("Vendor_Name", ""),
            "Vendor_CIN_No": st.session_state.answers.get("Vendor_CIN_No", ""),
            "Vendor_Office": st.session_state.answers.get("Vendor_Office", ""),
            "RPSG_Email": st.session_state.answers.get("RPSG_Email", ""),
            "Vendor_SPOC": st.session_state.answers.get("Vendor_SPOC", ""),
            "Vendor_Email": st.session_state.answers.get("Vendor_Email", ""),
            "RPSG_SPOC": st.session_state.answers.get("RPSG_SPOC", "")
        }
    elif contract_type == "PURCHASE AGREEMENT":
        placeholder_mapping = {
            key: st.session_state.answers.get(key, "")
            for key in st.session_state.answers.keys()
        }
    else:
        placeholder_mapping = {}

    if contract_type == "NDA":
        data = {k: st.session_state.answers.get(k, "") for k in [
            "DAY","MONTH","RPSG_CIN_No","Vendor_Name","Vendor_CIN_No","Vendor_Office",
            "RPSG_Email","Vendor_SPOC","Vendor_Email","RPSG_SPOC"
        ]}
        preamble_html = NDA_TEMPLATE.format(**data)
        title = "NON-DISCLOSURE AGREEMENT"
    elif contract_type == "PURCHASE AGREEMENT":
        data = {k: st.session_state.answers.get(k, "") for k in st.session_state.answers.keys()}
        preamble_html = PURCHASE_TEMPLATE
        title = "PURCHASE AGREEMENT"
    else:
        preamble_html = f"<p>{contract_type} between parties.</p>"
        title = contract_type or "AGREEMENT"

    # Build a unified list of sections for preview (standard + custom)
    std_sections = []
    if contract_type in ("NDA", "PURCHASE AGREEMENT"):
        for k in sorted(st.session_state.clause_store.keys()):
            std_sections.append((k, extract_title_from_clause(st.session_state.clause_store[k]), st.session_state.clause_store[k]))

    start_num = len(std_sections) + 1
    custom_sections = [(start_num + idx - 1, c.get("title","Custom"), c.get("text","")) 
                    for idx, c in enumerate(st.session_state.custom_clauses, 1)]
    
    # def highlight_html(text, terms):
    #     if not text or not terms:
    #         return text or ""
    #     rx = _regex_parts_for_terms(terms)
    #     def repl(m):
    #         return f"<u><mark>{m.group(0)}</mark></u>"
    #     return rx.sub(repl, text)

    st.markdown("### 📑 Preview")
    st.markdown(f"<h3 style='text-align:center'>{title}</h3>", unsafe_allow_html=True)

    #st.markdown(highlight_html(preamble_html, highlight_terms), unsafe_allow_html=True)
    st.markdown(highlight_placeholders(preamble_html, placeholder_mapping), unsafe_allow_html=True)

    st.markdown("<hr/>", unsafe_allow_html=True)

    if std_sections:
        st.markdown("**Sections**", unsafe_allow_html=True)
        for k, t, body in std_sections:
            st.markdown(f"**{t}**", unsafe_allow_html=True)
            st.markdown(highlight_placeholders(body, placeholder_mapping), unsafe_allow_html=True)
            st.markdown("<br/>", unsafe_allow_html=True)

    if custom_sections:
        st.markdown("**Custom Sections**")
        for num, t, body in custom_sections:
            st.markdown(f"**{num}. {t}**", unsafe_allow_html=True)
            st.markdown(highlight_placeholders(body, placeholder_mapping), unsafe_allow_html=True)
            st.markdown("<br/>", unsafe_allow_html=True)

    if st.session_state.parties:
        st.markdown("**PARTIES**")
        st.markdown("<ul>" + "".join([f"<li>{highlight_placeholders(p, placeholder_mapping)}</li>" for p in st.session_state.parties]) + "</ul>", unsafe_allow_html=True)

    #if contract_type != "NDA" and st.session_state.annexures:
    if st.session_state.annexures:
    #so now annexure functionality will also be there for NDA.
        st.markdown("### Annexures")
        for annex in st.session_state.annexures:
            st.markdown(f"#### {annex.get('title', 'Annexure')}")
            if annex.get("note"):
                st.markdown(highlight_placeholders(annex["note"], placeholder_mapping), unsafe_allow_html=True)

            if annex.get("excel_data"):
                st.markdown(f"**Data Table (from Excel)**")
                try:
                    annex.get("excel_data").seek(0)
                    annex_df = pd.read_excel(annex["excel_data"])
                    st.dataframe(annex_df, use_container_width=True)
                except Exception as e:
                    st.warning(f"Could not read Excel annexure '{annex.get('title')}': {e}")
            
            word_data = annex.get("word_data", [])
            if word_data and len(word_data) > 0:
                st.markdown(f"**Data Table (from Word-style Table)**")
                header = word_data[0] if word_data else []
                body = word_data[1:] if len(word_data) > 1 else []
                render_table(header, body)


    # Downloads
    c1, c2, c3 = st.columns(3)
    with c1:
        docx_file = build_docx(
            preamble_html=preamble_html,
            clauses_dict={k: v for k, v in st.session_state.clause_store.items()},
            custom_clauses=st.session_state.custom_clauses,
            parties=st.session_state.parties,
            #annexures=st.session_state.annexures if contract_type != "NDA" else [],
            annexures=st.session_state.annexures, #so now annexure functionality will also be there for NDA.
            title=title,
            highlight_terms=None
            #So now no highlighting will happen in the docx. It will only be there in the preview. 
        )
        st.download_button(
            label="📥 Download (.docx)",
            data=docx_file,
            file_name=f"{title.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with c2:
        if st.button("👁️ Refresh Preview", use_container_width=True):
            st.rerun()
    with c3:
        if st.button("✅ Submit", use_container_width=True):
            st.success("Submitted! (Wire this to email/DMS as needed.)")

    b1, b2 = st.columns(2)
    if b1.button("⬅️ Back to Sections"):
        st.session_state.workflow = "clauses"
        st.rerun()
    if contract_type != "NDA" and b2.button("⬅️ Back to Annexure"):
        st.session_state.workflow = "annexure"
        st.rerun()